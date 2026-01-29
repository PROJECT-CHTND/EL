"""Document Parser for various file formats."""

from __future__ import annotations

import io
import json
import logging
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import BinaryIO

from el_core.schemas import DocumentChunk

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types."""
    
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    TXT = "txt"
    MD = "md"
    JSON = "json"
    UNKNOWN = "unknown"

@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    
    content: str
    document_type: DocumentType
    page_count: int = 1
    metadata: dict[str, str] = field(default_factory=dict)


class DocumentParser:
    """Parser for various document formats."""
    
    SUPPORTED_EXTENSIONS = {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".xlsx": DocumentType.XLSX,
        ".txt": DocumentType.TXT,
        ".md": DocumentType.MD,
        ".json": DocumentType.JSON,
    }
    
    MIME_TYPE_MAP = {
        "application/pdf": DocumentType.PDF,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": DocumentType.XLSX,
        "text/plain": DocumentType.TXT,
        "text/markdown": DocumentType.MD,
        "application/json": DocumentType.JSON,
    }
    
    @classmethod
    def detect_type(cls, filename: str, content_type: str | None = None) -> DocumentType:
        """Detect document type from filename or content type.
        
        Args:
            filename: The filename with extension.
            content_type: Optional MIME type.
            
        Returns:
            Detected document type.
        """
        # Try extension first
        ext = Path(filename).suffix.lower()
        if ext in cls.SUPPORTED_EXTENSIONS:
            return cls.SUPPORTED_EXTENSIONS[ext]
        
        # Try MIME type
        if content_type and content_type in cls.MIME_TYPE_MAP:
            return cls.MIME_TYPE_MAP[content_type]
        
        return DocumentType.UNKNOWN
    
    @classmethod
    def parse(
        cls,
        file: BinaryIO | bytes,
        filename: str,
        content_type: str | None = None,
    ) -> ParsedDocument:
        """Parse a document and extract text content.
        
        Args:
            file: File object or bytes content.
            filename: Original filename.
            content_type: Optional MIME type.
            
        Returns:
            ParsedDocument with extracted content.
            
        Raises:
            ValueError: If document type is not supported.
        """
        doc_type = cls.detect_type(filename, content_type)
        
        if doc_type == DocumentType.UNKNOWN:
            raise ValueError(f"Unsupported document type: {filename}")
        
        # Convert to bytes if needed
        if hasattr(file, "read"):
            content_bytes = file.read()
        else:
            content_bytes = file
        
        # Parse based on type
        parser_map = {
            DocumentType.PDF: cls._parse_pdf,
            DocumentType.DOCX: cls._parse_docx,
            DocumentType.XLSX: cls._parse_xlsx,
            DocumentType.TXT: cls._parse_text,
            DocumentType.MD: cls._parse_text,
            DocumentType.JSON: cls._parse_json,
        }
        
        parser = parser_map.get(doc_type)
        if parser is None:
            raise ValueError(f"No parser available for: {doc_type}")
        
        return parser(content_bytes, filename)
    
    @classmethod
    def _parse_pdf(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF document."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
            
            metadata = {}
            if reader.metadata:
                if reader.metadata.title:
                    metadata["title"] = reader.metadata.title
                if reader.metadata.author:
                    metadata["author"] = reader.metadata.author
            
            return ParsedDocument(
                content="\n\n".join(text_parts),
                document_type=DocumentType.PDF,
                page_count=len(reader.pages),
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to parse PDF {filename}: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")
    
    @classmethod
    def _parse_docx(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX document."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    table_rows.append(row_text)
                if table_rows:
                    text_parts.append("\n[Table]\n" + "\n".join(table_rows))
            
            metadata = {}
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            
            return ParsedDocument(
                content="\n\n".join(text_parts),
                document_type=DocumentType.DOCX,
                page_count=1,  # DOCX doesn't have clear page boundaries
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX {filename}: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    @classmethod
    def _parse_xlsx(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse XLSX spreadsheet."""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            
            text_parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_content = [f"=== Sheet: {sheet_name} ==="]
                
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(v.strip() for v in row_values):
                        sheet_content.append(" | ".join(row_values))
                
                if len(sheet_content) > 1:
                    text_parts.append("\n".join(sheet_content))
            
            wb.close()
            
            return ParsedDocument(
                content="\n\n".join(text_parts),
                document_type=DocumentType.XLSX,
                page_count=len(wb.sheetnames),
                metadata={"sheets": str(wb.sheetnames)},
            )
        except Exception as e:
            logger.error(f"Failed to parse XLSX {filename}: {e}")
            raise ValueError(f"Failed to parse XLSX: {e}")
    
    @classmethod
    def _parse_text(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse plain text or markdown."""
        try:
            # Try UTF-8 first, then fallback to other encodings
            for encoding in ["utf-8", "utf-8-sig", "cp932", "shift_jis", "euc-jp"]:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = content.decode("utf-8", errors="replace")
            
            doc_type = DocumentType.MD if filename.lower().endswith(".md") else DocumentType.TXT
            
            return ParsedDocument(
                content=text,
                document_type=doc_type,
                page_count=1,
            )
        except Exception as e:
            logger.error(f"Failed to parse text {filename}: {e}")
            raise ValueError(f"Failed to parse text: {e}")
    
    @classmethod
    def _parse_json(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse JSON document."""
        try:
            text = content.decode("utf-8")
            data = json.loads(text)
            
            # Format JSON nicely for LLM consumption
            formatted = json.dumps(data, ensure_ascii=False, indent=2)
            
            return ParsedDocument(
                content=formatted,
                document_type=DocumentType.JSON,
                page_count=1,
                metadata={"type": type(data).__name__},
            )
        except Exception as e:
            logger.error(f"Failed to parse JSON {filename}: {e}")
            raise ValueError(f"Failed to parse JSON: {e}")


def truncate_content(
    content: str, 
    max_chars: int = 100000,
    keep_start_ratio: float = 0.2,
) -> str:
    """Truncate content to fit within LLM context limits.
    
    For chronological documents (like diaries), keeps both the beginning
    for context and the end for recent information.
    
    Args:
        content: The content to truncate.
        max_chars: Maximum character count (default: 100,000 for GPT-5.2).
        keep_start_ratio: Ratio of start content to keep when truncating (default: 0.2 = 20%).
        
    Returns:
        Truncated content with indicator if truncated.
    """
    if len(content) <= max_chars:
        return content
    
    # Split allocation: 20% for start, 80% for end (prioritize recent content)
    start_chars = int(max_chars * keep_start_ratio)
    end_chars = max_chars - start_chars - 100  # Reserve 100 chars for separator
    
    # Get start portion
    start_portion = content[:start_chars]
    # Try to end at a paragraph boundary
    last_break = max(start_portion.rfind("\n\n"), start_portion.rfind("。\n"))
    if last_break > start_chars * 0.7:
        start_portion = start_portion[:last_break + 1]
    
    # Get end portion
    end_portion = content[-end_chars:]
    # Try to start at a paragraph boundary
    first_break = end_portion.find("\n\n")
    if first_break == -1:
        first_break = end_portion.find("\n")
    if first_break != -1 and first_break < end_chars * 0.3:
        end_portion = end_portion[first_break + 1:]
    
    omitted_chars = len(content) - len(start_portion) - len(end_portion)
    separator = f"\n\n[... 中間部分 約{omitted_chars:,}文字を省略（最新の内容を優先） ...]\n\n"
    
    return start_portion + separator + end_portion


# =============================================================================
# Document Chunking - Date-based Semantic Chunking
# =============================================================================

# Date patterns for chunking
# Priority order: more specific patterns first
DATE_PATTERNS: list[tuple[str, str]] = [
    # ISO format: 2024-01-15, 2024/01/15
    (r"(?:^|\n)#+\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})", "%Y-%m-%d"),
    (r"(?:^|\n)(\d{4}[-/]\d{1,2}[-/]\d{1,2})(?:\s|$|:)", "%Y-%m-%d"),
    # Japanese format: 2024年1月15日
    (r"(?:^|\n)#+\s*(\d{4}年\d{1,2}月\d{1,2}日)", "%Y年%m月%d日"),
    (r"(?:^|\n)(\d{4}年\d{1,2}月\d{1,2}日)", "%Y年%m月%d日"),
    # Short format with header: ## 1/15, ## 1月15日
    (r"(?:^|\n)#+\s*(\d{1,2}月\d{1,2}日)", "%m月%d日"),
    (r"(?:^|\n)#+\s*(\d{1,2}/\d{1,2})(?:\s|$)", "%m/%d"),
    # Diary format: 1/21(水), 12/19（金） - month/day followed by weekday in parentheses
    # This is a common Japanese diary format
    (r"(?:^|\n)(\d{1,2}/\d{1,2})[\(（]", "%m/%d"),
    # Date only (must be at line start to avoid false positives)
    (r"^(\d{1,2}月\d{1,2}日)", "%m月%d日"),
    (r"^(\d{1,2}/\d{1,2})(?:\s|:|\(|（)", "%m/%d"),
]


@dataclass
class DateMatch:
    """A date found in the document."""
    
    date: datetime
    position: int  # Character position in the document
    match_text: str  # The matched text
    line_start: int  # Start of the line containing the date


def extract_dates_from_content(
    content: str,
    base_year: int | None = None,
) -> list[DateMatch]:
    """Extract dates from document content.
    
    Handles year-spanning documents (e.g., December to January) by detecting
    when dates go backwards (Dec -> Jan) and adjusting years accordingly.
    
    Args:
        content: The document content.
        base_year: Year to use for dates without year (default: current year).
        
    Returns:
        List of DateMatch objects sorted by position.
    """
    if base_year is None:
        base_year = datetime.now().year
    
    matches: list[DateMatch] = []
    seen_positions: set[int] = set()
    
    for pattern, date_format in DATE_PATTERNS:
        for match in re.finditer(pattern, content, re.MULTILINE):
            date_str = match.group(1)
            pos = match.start(1)
            
            # Avoid duplicate matches at the same position
            if pos in seen_positions:
                continue
            
            try:
                # Parse the date
                parsed_date = datetime.strptime(date_str, date_format)
                
                # If no year in format, use base_year initially
                if "%Y" not in date_format:
                    parsed_date = parsed_date.replace(year=base_year)
                
                # Find the start of the line containing this date
                line_start = content.rfind("\n", 0, pos)
                line_start = line_start + 1 if line_start != -1 else 0
                
                matches.append(DateMatch(
                    date=parsed_date,
                    position=pos,
                    match_text=date_str,
                    line_start=line_start,
                ))
                seen_positions.add(pos)
            except ValueError:
                # Invalid date, skip
                continue
    
    # Sort by position in document
    matches.sort(key=lambda m: m.position)
    
    # Fix year-spanning documents (e.g., Dec -> Jan transition)
    # If we see dates going backwards (Dec -> Jan), adjust earlier dates to previous year
    if len(matches) >= 2:
        matches = _fix_year_spanning_dates(matches, base_year)
    
    return matches


def _fix_year_spanning_dates(matches: list[DateMatch], base_year: int) -> list[DateMatch]:
    """Fix dates in year-spanning documents.
    
    When a document spans from December to January (or similar year transitions),
    the earlier dates need to be adjusted to the previous year.
    
    Args:
        matches: List of DateMatch objects sorted by position.
        base_year: The base year used for parsing.
        
    Returns:
        Adjusted list of DateMatch objects.
    """
    # Detect year transitions by looking for month decreases (e.g., Dec to Jan)
    # Going from month 12 to month 1 indicates a year change
    
    adjusted_matches: list[DateMatch] = []
    year_offset = 0  # How many years to subtract from base_year
    prev_month = None
    
    for match in matches:
        current_month = match.date.month
        
        # Check for year transition (month going backwards significantly)
        # Dec (12) -> Jan (1) or any large backward jump
        if prev_month is not None and prev_month > current_month and prev_month - current_month > 6:
            # We crossed into a new year, previous dates were from last year
            # But actually, we need to KEEP current dates as base_year
            # and ADJUST all previous dates to be last year
            pass  # Will handle below
        
        prev_month = current_month
    
    # Alternative approach: if first date is Dec and later dates are Jan,
    # the Dec dates should be previous year
    first_month = matches[0].date.month
    
    # Find if there's a Dec -> Jan transition
    has_transition = False
    for i, match in enumerate(matches):
        if i > 0:
            prev_match = matches[i - 1]
            if prev_match.date.month == 12 and match.date.month == 1:
                has_transition = True
                break
            # Also handle 11->1, 10->1 etc for documents starting mid-year
            if prev_match.date.month > match.date.month and prev_match.date.month - match.date.month > 6:
                has_transition = True
                break
    
    if has_transition:
        # Adjust: dates before the transition are from previous year
        in_new_year = False
        for i, match in enumerate(matches):
            if i > 0:
                prev_match = adjusted_matches[-1] if adjusted_matches else matches[i - 1]
                orig_prev = matches[i - 1]
                # Detect transition
                if orig_prev.date.month > match.date.month and orig_prev.date.month - match.date.month > 6:
                    in_new_year = True
            
            if not in_new_year:
                # Adjust to previous year
                new_date = match.date.replace(year=base_year - 1)
                adjusted_matches.append(DateMatch(
                    date=new_date,
                    position=match.position,
                    match_text=match.match_text,
                    line_start=match.line_start,
                ))
            else:
                # Keep current year
                adjusted_matches.append(match)
        
        return adjusted_matches
    
    return matches


def chunk_by_date(
    content: str,
    document_id: str,
    base_year: int | None = None,
    min_chunk_size: int = 10,
) -> list[DocumentChunk]:
    """Split document content into chunks by date.
    
    Each chunk represents a date entry (e.g., a diary entry for one day).
    The original content is preserved exactly for accurate reference.
    
    Args:
        content: The document content.
        document_id: ID of the parent document.
        base_year: Year to use for dates without year (default: current year).
        min_chunk_size: Minimum characters for a valid chunk.
        
    Returns:
        List of DocumentChunk objects.
    """
    date_matches = extract_dates_from_content(content, base_year)
    
    chunks: list[DocumentChunk] = []
    
    if not date_matches:
        # No dates found - create single chunk or split by paragraphs
        return chunk_by_paragraphs(content, document_id)
    
    for i, date_match in enumerate(date_matches):
        # Chunk starts at the line containing the date
        chunk_start = date_match.line_start
        
        # Chunk ends at the start of the next date's line, or end of document
        if i + 1 < len(date_matches):
            chunk_end = date_matches[i + 1].line_start
        else:
            chunk_end = len(content)
        
        chunk_content = content[chunk_start:chunk_end].strip()
        
        # Skip very small chunks
        if len(chunk_content) < min_chunk_size:
            continue
        
        # Extract heading (first line if it looks like a heading)
        lines = chunk_content.split("\n", 1)
        heading = ""
        if lines[0].startswith("#") or re.match(r"^\d", lines[0]):
            heading = lines[0].strip("# \t")
        
        chunk = DocumentChunk(
            id=str(uuid.uuid4()),
            document_id=document_id,
            content=chunk_content,
            chunk_index=len(chunks),
            chunk_date=date_match.date,
            heading=heading,
            char_count=len(chunk_content),
            created_at=datetime.now(),
        )
        chunks.append(chunk)
    
    return chunks


def chunk_by_paragraphs(
    content: str,
    document_id: str,
    max_chunk_size: int = 2000,
    min_chunk_size: int = 50,
) -> list[DocumentChunk]:
    """Split document content into chunks by paragraphs.
    
    Used as fallback when no dates are detected.
    
    Args:
        content: The document content.
        document_id: ID of the parent document.
        max_chunk_size: Maximum characters per chunk.
        min_chunk_size: Minimum characters for a valid chunk.
        
    Returns:
        List of DocumentChunk objects.
    """
    # Split by double newlines (paragraph boundaries)
    paragraphs = re.split(r"\n\n+", content)
    
    chunks: list[DocumentChunk] = []
    current_chunk_parts: list[str] = []
    current_size = 0
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        para_size = len(para)
        
        # If adding this paragraph would exceed max size, save current chunk
        if current_size + para_size > max_chunk_size and current_chunk_parts:
            chunk_content = "\n\n".join(current_chunk_parts)
            if len(chunk_content) >= min_chunk_size:
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    content=chunk_content,
                    chunk_index=len(chunks),
                    chunk_date=None,
                    heading=_extract_heading(current_chunk_parts[0]),
                    char_count=len(chunk_content),
                    created_at=datetime.now(),
                )
                chunks.append(chunk)
            current_chunk_parts = []
            current_size = 0
        
        current_chunk_parts.append(para)
        current_size += para_size
    
    # Don't forget the last chunk
    if current_chunk_parts:
        chunk_content = "\n\n".join(current_chunk_parts)
        if len(chunk_content) >= min_chunk_size:
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                document_id=document_id,
                content=chunk_content,
                chunk_index=len(chunks),
                chunk_date=None,
                heading=_extract_heading(current_chunk_parts[0]),
                char_count=len(chunk_content),
                created_at=datetime.now(),
            )
            chunks.append(chunk)
    
    # If no chunks created, create one for the entire content
    if not chunks and content.strip():
        chunk = DocumentChunk(
            id=str(uuid.uuid4()),
            document_id=document_id,
            content=content.strip(),
            chunk_index=0,
            chunk_date=None,
            heading="",
            char_count=len(content.strip()),
            created_at=datetime.now(),
        )
        chunks.append(chunk)
    
    return chunks


def _extract_heading(text: str) -> str:
    """Extract heading from the first line of text."""
    first_line = text.split("\n", 1)[0].strip()
    if first_line.startswith("#"):
        return first_line.strip("# \t")
    elif len(first_line) <= 50:  # Short lines might be headings
        return first_line
    return ""


def chunk_document(
    parsed_doc: ParsedDocument,
    document_id: str,
    base_year: int | None = None,
) -> list[DocumentChunk]:
    """Chunk a parsed document using appropriate strategy.
    
    Automatically selects between date-based and paragraph-based chunking.
    
    Args:
        parsed_doc: The parsed document.
        document_id: ID of the parent document.
        base_year: Year to use for dates without year.
        
    Returns:
        List of DocumentChunk objects.
    """
    content = parsed_doc.content
    
    # Try date-based chunking first
    chunks = chunk_by_date(content, document_id, base_year)
    
    # If date-based chunking resulted in chunks, use those
    # Otherwise, the function already fell back to paragraph-based
    
    logger.info(f"Chunked document {document_id} into {len(chunks)} chunks")
    
    return chunks
